from docx import Document

doc = Document();

doc.add_heading("Hello world")

p = doc.add_paragraph("sample text")
p.add_run(" bold test").bold = True;

doc.save("example.docx")